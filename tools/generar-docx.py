#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el documento Word corporativo de copias de seguridad de OpenCode.

CODANOR - Jafa, S.L. | Barcelona, Catalunya

Construye 'documento-copias-seguridad.docx' aplicando la identidad visual
corporativa (colores, tipografias y estructura) a partir del mismo contenido
que 'documento-copias-seguridad.html'.

Uso:
    python3 tools/generar-docx.py

Requisitos:
    pip install python-docx
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# --- Paleta corporativa CODANOR ---------------------------------------------
PRIMARY = RGBColor(0x17, 0x8D, 0xC2)
PRIMARY_DARK = RGBColor(0x12, 0x75, 0xA3)
ACCENT = RGBColor(0x12, 0xA7, 0x9D)
DARK = RGBColor(0x1F, 0x29, 0x33)
GRAY = RGBColor(0x52, 0x60, 0x6D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

# Tonos de fondo para alertas y cabeceras (en formato hex sin '#')
FILL_PRIMARY = '178DC2'
FILL_PRIMARY_LIGHT = 'E8F4FA'
FILL_ACCENT_LIGHT = 'E6F6F5'
FILL_WARNING_LIGHT = 'FFF6E0'
FILL_DANGER_LIGHT = 'FCEAEA'
FILL_CODE = '1F2933'
FILL_CODE_TEXT = 'E4E7EB'

# Tipografias (con fallback automatico de Word si no estan instaladas)
FONT_HEADING = 'Poppins'
FONT_BODY = 'Nunito Sans'
FONT_MONO = 'Consolas'

OUTPUT = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    'documento-copias-seguridad.docx',
)


# --- Utilidades de bajo nivel (OOXML) ---------------------------------------
def set_cell_background(cell, fill_hex):
  """Aplica un color de fondo solido a una celda de tabla."""
  tc_pr = cell._tc.get_or_add_tcPr()
  shd = OxmlElement('w:shd')
  shd.set(qn('w:val'), 'clear')
  shd.set(qn('w:color'), 'auto')
  shd.set(qn('w:fill'), fill_hex)
  tc_pr.append(shd)


def set_paragraph_background(paragraph, fill_hex):
  """Aplica un sombreado de fondo a un parrafo completo."""
  p_pr = paragraph._p.get_or_add_pPr()
  shd = OxmlElement('w:shd')
  shd.set(qn('w:val'), 'clear')
  shd.set(qn('w:color'), 'auto')
  shd.set(qn('w:fill'), fill_hex)
  p_pr.append(shd)


def set_left_border(paragraph, color_hex):
  """Anade un borde izquierdo de color al parrafo (estilo 'alerta')."""
  p_pr = paragraph._p.get_or_add_pPr()
  borders = OxmlElement('w:pBdr')
  left = OxmlElement('w:left')
  left.set(qn('w:val'), 'single')
  left.set(qn('w:sz'), '24')
  left.set(qn('w:space'), '8')
  left.set(qn('w:color'), color_hex)
  borders.append(left)
  p_pr.append(borders)


# --- Constructores de contenido ---------------------------------------------
def add_run(paragraph, text, *, bold=False, color=None, font=FONT_BODY,
            size=11, mono=False):
  run = paragraph.add_run(text)
  run.bold = bold
  run.font.name = FONT_MONO if mono else font
  run.font.size = Pt(size)
  if color is not None:
    run.font.color.rgb = color
  return run


def add_heading(doc, text, level):
  """Titulo de seccion con color corporativo."""
  p = doc.add_paragraph()
  p.space_before = Pt(12)
  if level == 1:
    add_run(p, text, bold=True, color=PRIMARY_DARK, font=FONT_HEADING, size=18)
  elif level == 2:
    add_run(p, text, bold=True, color=PRIMARY, font=FONT_HEADING, size=14)
  else:
    add_run(p, text, bold=True, color=ACCENT, font=FONT_HEADING, size=12)
  return p


def add_body(doc, text):
  p = doc.add_paragraph()
  add_run(p, text)
  return p


def add_bullets(doc, items):
  for item in items:
    p = doc.add_paragraph(style='List Bullet')
    # Permite resaltar la parte previa a ':' en negrita.
    if ': ' in item:
      head, tail = item.split(': ', 1)
      add_run(p, head + ': ', bold=True)
      add_run(p, tail)
    else:
      add_run(p, item)


def add_steps(doc, items):
  for item in items:
    p = doc.add_paragraph(style='List Number')
    if ': ' in item:
      head, tail = item.split(': ', 1)
      add_run(p, head + ': ', bold=True)
      add_run(p, tail)
    else:
      add_run(p, item)


def add_code(doc, code):
  """Bloque de codigo con fondo oscuro y fuente monoespaciada."""
  p = doc.add_paragraph()
  set_paragraph_background(p, FILL_CODE)
  for i, line in enumerate(code.strip('\n').split('\n')):
    if i > 0:
      p.add_run().add_break()
    run = p.add_run(line)
    run.font.name = FONT_MONO
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0xE4, 0xE7, 0xEB)


def add_alert(doc, title, body, kind='info'):
  """Bloque de alerta con titulo, borde de color y fondo suave."""
  palette = {
    'info': (PRIMARY_DARK, FILL_PRIMARY_LIGHT, '178DC2'),
    'success': (ACCENT, FILL_ACCENT_LIGHT, '12A79D'),
    'warning': (RGBColor(0xB6, 0x7A, 0x00), FILL_WARNING_LIGHT, 'F0A202'),
    'danger': (RGBColor(0xD6, 0x45, 0x45), FILL_DANGER_LIGHT, 'D64545'),
  }
  title_color, fill, border = palette[kind]

  p_title = doc.add_paragraph()
  set_paragraph_background(p_title, fill)
  set_left_border(p_title, border)
  add_run(p_title, title, bold=True, color=title_color, font=FONT_HEADING,
          size=11)

  p_body = doc.add_paragraph()
  set_paragraph_background(p_body, fill)
  set_left_border(p_body, border)
  add_run(p_body, body, color=DARK)


def add_table(doc, headers, rows):
  """Tabla con cabecera en color primary y filas de datos."""
  table = doc.add_table(rows=1, cols=len(headers))
  table.alignment = WD_TABLE_ALIGNMENT.CENTER
  table.style = 'Table Grid'

  hdr_cells = table.rows[0].cells
  for i, text in enumerate(headers):
    set_cell_background(hdr_cells[i], FILL_PRIMARY)
    para = hdr_cells[i].paragraphs[0]
    add_run(para, text, bold=True, color=WHITE, font=FONT_HEADING, size=10)

  for row in rows:
    cells = table.add_row().cells
    for i, text in enumerate(row):
      para = cells[i].paragraphs[0]
      add_run(para, text, size=10)
  doc.add_paragraph()
  return table


# --- Portada ----------------------------------------------------------------
def build_cover(doc):
  for _ in range(6):
    doc.add_paragraph()

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'CODANOR', bold=True, color=PRIMARY, font=FONT_HEADING, size=28)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'Copias de Seguridad de OpenCode', bold=True, color=DARK,
          font=FONT_HEADING, size=26)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'Opciones de respaldo, automatizacion y recuperacion',
          color=GRAY, font=FONT_HEADING, size=14)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p,
          'Guia corporativa sobre las posibilidades de realizar copias de '
          'seguridad de proyectos, datos y configuracion de OpenCode, con '
          'enfasis en el almacenamiento en NAS Synology.',
          color=GRAY)

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'Version 1.0  -  Junio 2026', bold=True, color=ACCENT,
          font=FONT_HEADING, size=11)

  for _ in range(8):
    doc.add_paragraph()

  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'Jafa, S.L. - CODANOR - Barcelona, Catalunya', color=GRAY,
          size=10)

  doc.add_page_break()


# --- Documento completo -----------------------------------------------------
def build_document():
  doc = Document()

  # Estilo base del cuerpo.
  normal = doc.styles['Normal']
  normal.font.name = FONT_BODY
  normal.font.size = Pt(11)
  normal.font.color.rgb = DARK

  build_cover(doc)

  # 1. Introduccion
  add_heading(doc, '1. Introduccion', 1)
  add_body(doc,
           'El entorno de trabajo con OpenCode en CODANOR genera tres activos '
           'digitales que deben protegerse frente a perdida de datos, fallos '
           'de hardware o errores humanos:')
  add_bullets(doc, [
    'Proyectos de codigo: el trabajo desarrollado en ~/Proyectos.',
    'Datos de OpenCode: el historial de sesiones, almacenamiento interno y la '
    'base de datos opencode.db.',
    'Configuracion: agentes, comandos y ajustes en ~/.config/opencode.',
  ])
  add_body(doc,
           'Este documento describe las opciones disponibles para realizar '
           'dichas copias de forma fiable y recurrente, haciendo especial '
           'enfasis en el uso de un NAS Synology como destino principal de '
           'respaldo.')
  add_alert(doc, 'Regla 3-2-1',
            'Mantener 3 copias de los datos, en 2 tipos de soporte distintos, '
            'con 1 copia fuera del equipo principal. El NAS Synology cubre la '
            'copia centralizada y duradera de esta estrategia.', 'info')

  # 2. Que se respalda y que se excluye
  add_heading(doc, '2. Que se respalda y que se excluye', 1)
  add_body(doc,
           'Para optimizar espacio y tiempo, el respaldo incluye solo los '
           'datos imprescindibles y excluye material voluminoso y reproducible.')
  add_heading(doc, 'Incluido en la copia', 2)
  add_table(doc, ['Bloque', 'Origen', 'Tamano aprox.'], [
    ['Proyectos', '~/Proyectos (18 proyectos)', '~5.6 GB'],
    ['Datos OpenCode', 'storage/ + opencode.db + auth.json + account.json',
     '~640 MB'],
    ['Configuracion', '~/.config/opencode', '~6.5 MB'],
  ])
  add_heading(doc, 'Excluido de la copia', 2)
  add_table(doc, ['Elemento', 'Motivo de exclusion', 'Tamano evitado'], [
    ['snapshot/', 'Voluminoso y regenerable', '~19 GB'],
    ['cache/', 'Datos temporales reproducibles', '~102 MB'],
    ['log/', 'Registros no criticos', 'variable'],
    ['bin/', 'Binarios reinstalables', 'variable'],
  ])
  add_alert(doc, 'Resultado',
            'Al excluir los snapshots, el tamano total del respaldo se reduce '
            'de unos 26 GB a aproximadamente 6.3 GB, agilizando cada copia y '
            'reduciendo el consumo de almacenamiento en el NAS.', 'success')

  # 3. Opciones de realizacion del backup
  add_heading(doc, '3. Opciones de realizacion del backup', 1)
  add_body(doc,
           'Existen varias vias para almacenar las copias. La siguiente '
           'comparativa resume las disponibles en el entorno CODANOR.')
  add_table(doc, ['Opcion', 'Tipo', 'Descripcion', 'Ideal para'], [
    ['NAS Synology (RECOMENDADO)', 'Principal',
     'Almacenamiento en red centralizado, RAID, 7.4 TB libres, retencion '
     'prolongada.', 'Copia principal automatizada'],
    ['Time Machine', 'Sistema',
     'Copia horaria del sistema completo de macOS (disco 1.8 TB).',
     'Recuperacion del sistema'],
    ['Disco externo USB', 'Local',
     'Copia en unidad externa (p. ej. Transcend, 852 GB).',
     'Copia offline portatil'],
    ['Nube / GitHub', 'Futuro',
     'Tercera capa fuera de las instalaciones (por definir).',
     'Copia externa (off-site)'],
  ])

  # 4. NAS Synology
  add_heading(doc, '4. Copia de seguridad en NAS Synology', 1)
  add_body(doc,
           'El NAS Synology es la opcion recomendada como destino principal de '
           'las copias de OpenCode en CODANOR. Un NAS (Network Attached '
           'Storage) es un sistema de almacenamiento conectado a la red local '
           'que ofrece capacidad, redundancia y acceso compartido.')
  add_heading(doc, 'Ventajas frente a otras opciones', 2)
  add_bullets(doc, [
    'Centralizado: accesible desde cualquier equipo de la red.',
    'Redundancia RAID: tolerancia a fallo de disco.',
    'Gran capacidad: 7.4 TB libres en el volumen CODANOR.',
    'Retencion prolongada: permite conservar muchas versiones.',
    'Funciones DSM: snapshots, papelera de red, cuotas y Hyper Backup.',
  ])

  add_heading(doc, '4.1. Montar el NAS en macOS (SMB)', 2)
  add_body(doc, 'El volumen se conecta mediante el protocolo SMB. Desde Finder:')
  add_steps(doc, [
    'En Finder, menu Ir > Conectarse al servidor (Cmd+K).',
    'Introducir la direccion del NAS, por ejemplo '
    'smb://SYCODANOR-920.local/CODANOR.',
    'Autenticarse con el usuario del NAS (p. ej. Jfajeda).',
    'El volumen quedara montado en /Volumes/CODANOR.',
  ])
  add_alert(doc, 'Requisito imprescindible en macOS: Acceso total al disco',
            'Para que un script o tarea automatica pueda escribir en un volumen '
            'de red, macOS exige conceder el permiso "Acceso total al disco" '
            '(TCC). Sin el, los comandos devuelven "Operation not permitted". '
            'Concedelo en Ajustes del Sistema > Privacidad y seguridad > Acceso '
            'total al disco, anadiendo la aplicacion de Terminal (y /bin/bash si '
            'se usa launchd). Despues, reinicia la Terminal.', 'warning')

  add_heading(doc, '4.2. Configurar el script de backup hacia el NAS', 2)
  add_body(doc,
           'El script backup-opencode.sh tiene como destino por defecto el NAS. '
           'Tambien puede indicarse explicitamente mediante variables de '
           'entorno:')
  add_code(doc,
           'export OPENCODE_BACKUP_DIR="/Volumes/CODANOR/opencode-backups"\n'
           'export OPENCODE_MIN_FREE_MB=8000\n\n'
           '# Backup completo (proyectos + datos + configuracion)\n'
           './backup-opencode.sh\n\n'
           '# Solo datos de OpenCode (storage + DB consistente)\n'
           './backup-opencode.sh --sessions')
  add_body(doc,
           'La base de datos opencode.db se copia de forma consistente con '
           'sqlite3 ".backup", evitando corrupcion por el modo WAL. Antes de '
           'copiar, el script valida que el volumen este montado, realiza una '
           'prueba de escritura real y comprueba el espacio libre minimo.')

  add_heading(doc, '4.3. Automatizacion (copias recurrentes)', 2)
  add_body(doc,
           'Para respaldar de forma sucesiva y desatendida, se emplea launchd '
           '(recomendado en macOS):')
  add_code(doc,
           'cp com.codanor.opencode-backup.plist ~/Library/LaunchAgents/\n'
           'launchctl load ~/Library/LaunchAgents/'
           'com.codanor.opencode-backup.plist')
  add_body(doc,
           'El agente ejecuta el backup diariamente a las 02:00. Alternativa '
           'con cron:')
  add_code(doc,
           '0 2 * * * /ruta/al/backup-opencode.sh >> '
           '/Volumes/CODANOR/opencode-backups/backup.log 2>&1')

  add_heading(doc, '4.4. Buenas practicas en Synology DSM', 2)
  add_body(doc, 'Desde el panel DSM del NAS Synology conviene configurar:')
  add_table(doc, ['Funcion DSM', 'Recomendacion'], [
    ['Carpeta compartida',
     'Crear una carpeta dedicada (opencode-backups) con permisos restringidos '
     'al usuario de backup.'],
    ['Snapshot Replication',
     'Activar snapshots programados de la carpeta compartida (Btrfs) para '
     'versiones inmutables.'],
    ['Papelera de red',
     'Habilitar la papelera (#recycle) para recuperar archivos borrados por '
     'error.'],
    ['Cuotas de usuario',
     'Definir una cuota para evitar que el backup agote el volumen.'],
    ['Hyper Backup',
     'Replicar la carpeta a otro NAS o a la nube (copia off-site / regla '
     '3-2-1).'],
    ['SMB version',
     'Forzar SMB 3 y cifrado en transito desde Panel de control > Servicio de '
     'archivos.'],
  ])
  add_alert(doc, 'Estructura de archivos en el NAS',
            'Cada copia genera ficheros con marca temporal: '
            'proyectos_AAAAMMDD_HHMMSS.tar.gz, '
            'opencode_data_AAAAMMDD_HHMMSS.tar.gz y '
            'opencode_config_AAAAMMDD_HHMMSS.tar.gz, junto a un registro '
            'backup.log. Se conservan las ultimas 30 copias por tipo (rotacion '
            'automatica).', 'info')

  # 5. Estrategia recomendada
  add_heading(doc, '5. Estrategia recomendada', 1)
  add_body(doc,
           'La combinacion de capas garantiza resiliencia ante distintos tipos '
           'de fallo:')
  add_table(doc, ['Capa', 'Soporte', 'Frecuencia', 'Cubre'], [
    ['1 - Principal', 'NAS Synology', 'Diaria (02:00)',
     'Proyectos + datos + config'],
    ['2 - Sistema', 'Time Machine', 'Horaria', 'Sistema completo de macOS'],
    ['3 - Off-site (futuro)', 'Nube / Hyper Backup', 'Por definir',
     'Copia externa fuera de sede'],
  ])
  add_alert(doc, 'Recomendacion CODANOR',
            'Mantener el NAS Synology como destino principal automatizado y '
            'Time Machine como copia del sistema. Cuando se habilite la capa '
            'off-site (Hyper Backup a la nube), se completara la regla 3-2-1.',
            'success')

  # 6. Restauracion
  add_heading(doc, '6. Restauracion de copias', 1)
  add_body(doc, 'El propio script permite listar y restaurar copias existentes:')
  add_code(doc,
           '# Listar copias disponibles en el NAS\n'
           './backup-opencode.sh --list\n\n'
           '# Restaurar una copia concreta (crea un backup de seguridad '
           'previo)\n'
           './backup-opencode.sh --restore '
           'opencode_data_20260606_020000.tar.gz')
  add_body(doc,
           'Antes de sobrescribir, el proceso de restauracion genera '
           'automaticamente una copia de seguridad del estado actual, de modo '
           'que una restauracion erronea sea reversible.')
  add_body(doc, 'Tambien es posible restaurar manualmente un archivo comprimido:')
  add_code(doc,
           '# Verificar integridad del paquete antes de restaurar\n'
           'tar -tzf opencode_data_20260606_020000.tar.gz | head\n\n'
           '# Extraer en el directorio de usuario\n'
           'tar -xzf opencode_data_20260606_020000.tar.gz -C ~/\n\n'
           '# Comprobar la integridad de la base de datos restaurada\n'
           'sqlite3 ~/.local/share/opencode/opencode.db '
           '"PRAGMA integrity_check;"')

  # 7. Recuperacion ante desastre total
  add_heading(doc, '7. Recuperacion ante desastre total', 1)
  add_body(doc,
           'En caso de perdida completa del equipo (robo, averia irreparable o '
           'sustitucion por un Mac nuevo), el procedimiento de recuperacion es '
           'el siguiente:')
  add_steps(doc, [
    'Instalar OpenCode en el equipo nuevo.',
    'Montar el NAS Synology por SMB en /Volumes/CODANOR.',
    'Conceder "Acceso total al disco" a la Terminal en Ajustes del Sistema.',
    'Restaurar la configuracion: extraer el ultimo opencode_config_*.tar.gz en '
    '~/.',
    'Restaurar los datos: extraer el ultimo opencode_data_*.tar.gz (incluye '
    'sesiones, DB, credenciales).',
    'Restaurar los proyectos: extraer el ultimo proyectos_*.tar.gz en ~/.',
    'Verificar la integridad de la base de datos con PRAGMA integrity_check.',
    'Reactivar la automatizacion cargando de nuevo el agente launchd.',
  ])
  add_alert(doc, 'Importante',
            'El archivo auth.json contiene credenciales. Tratalo como '
            'informacion sensible: nunca lo subas a repositorios publicos y '
            'restaura los datos solo en equipos de confianza de CODANOR.',
            'danger')

  # Footer
  doc.add_paragraph()
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'CODANOR - Jafa, S.L.', bold=True, color=DARK, font=FONT_HEADING,
          size=10)
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p, 'Barcelona, Catalunya', color=GRAY, size=9)
  p = doc.add_paragraph()
  p.alignment = WD_ALIGN_PARAGRAPH.CENTER
  add_run(p,
          'Documento interno - Copias de Seguridad de OpenCode - Version 1.0 '
          '(Junio 2026)', color=GRAY, size=9)

  doc.save(OUTPUT)
  print('Documento generado: ' + OUTPUT)


if __name__ == '__main__':
  build_document()

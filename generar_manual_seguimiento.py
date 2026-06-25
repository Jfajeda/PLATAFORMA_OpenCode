#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera el Manual Word "Seguimiento de Reuniones con IA" para CODANOR.
Branding: paleta corporativa (#178DC2 primario, #12A79D acento) y logo de codanor.com.
Requiere: python-docx (pip install python-docx).
"""

import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- Branding CODANOR ---
PRIMARY = RGBColor(0x17, 0x8D, 0xC2)   # #178DC2
ACCENT = RGBColor(0x12, 0xA7, 0x9D)    # #12A79D
DARK = RGBColor(0x0D, 0x3B, 0x66)      # azul oscuro para texto destacado
GREY = RGBColor(0x55, 0x5B, 0x66)
LIGHT_BG = "EAF4FA"                     # fondo suave para cajas

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO = os.path.join(BASE_DIR, "Biblioteca", "codanor_logo.png")
OUT = os.path.join(BASE_DIR, "Manual_Seguimiento_Reuniones_IA_Codanor.docx")

doc = Document()

# --- Estilos base ---
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x22, 0x26, 0x2B)


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = PRIMARY
        p.space_before = Pt(14)
        p.space_after = Pt(6)
        # línea inferior
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "12A79D")
        pbdr.append(bottom)
        pPr.append(pbdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = DARK
        p.space_before = Pt(10)
        p.space_after = Pt(3)
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = ACCENT
        p.space_before = Pt(8)
        p.space_after = Pt(2)
    return p


def add_body(text, bold=False, italic=False, color=None, size=11, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = DARK
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_number(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Number")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.color.rgb = DARK
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_callout(title, text, fill=LIGHT_BG, title_color=PRIMARY):
    """Caja destacada con título y cuerpo."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, fill)
    p_title = cell.paragraphs[0]
    rt = p_title.add_run(title)
    rt.bold = True
    rt.font.size = Pt(11)
    rt.font.color.rgb = title_color
    p_body = cell.add_paragraph()
    rb = p_body.add_run(text)
    rb.font.size = Pt(10.5)
    # margen interior
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def add_steps_table(rows):
    """Tabla de pasos (nº, acción)."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Light Grid Accent 1"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Paso"
    hdr[1].text = "Acción"
    for c in hdr:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    for i, (paso, accion) in enumerate(rows, 1):
        cells = tbl.add_row().cells
        cells[0].text = str(i)
        cells[1].text = accion
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def set_footer():
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CODANOR · Jafa, S.L. — Plataforma de Seguimiento de Proyectos · Manual de uso interno")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ============================================================
# PORTADA
# ============================================================
section = doc.sections[0]
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)

# Logo
if os.path.exists(LOGO):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.add_run().add_picture(LOGO, width=Inches(3.2))

doc.add_paragraph().paragraph_format.space_after = Pt(30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Manual de Seguimiento de Reuniones con IA")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = PRIMARY

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Módulo de Consultoría — Pestaña «Seguimiento»")
r.font.size = Pt(15)
r.font.color.rgb = DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Plataforma de Seguimiento de Proyectos · ISO 27001 (referencia)")
r.font.size = Pt(12)
r.font.color.rgb = GREY

doc.add_paragraph().paragraph_format.space_after = Pt(40)

# Cuadro de datos del documento
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
datos = [
    ("Versión del documento", "v1.0"),
    ("Responsable", "Jafa, S.L. (CODANOR)"),
    ("Fecha", date.today().strftime("%d/%m/%Y")),
    ("Ámbito", "Uso interno — Consultoría Organizativa y Normativa"),
]
for i, (k, v) in enumerate(datos):
    c0, c1 = tbl.rows[i].cells
    c0.text = ""
    rk = c0.paragraphs[0].add_run(k)
    rk.bold = True
    rk.font.color.rgb = DARK
    rk.font.size = Pt(10.5)
    set_cell_shading(c0, LIGHT_BG)
    c1.text = ""
    rv = c1.paragraphs[0].add_run(v)
    rv.font.size = Pt(10.5)

doc.add_page_break()

# ============================================================
# 1. INTRODUCCIÓN
# ============================================================
add_heading("1. Introducción", 1)
add_body(
    "La pestaña «Seguimiento», dentro del módulo de Consultoría (Implantación del SGSI), "
    "permite gestionar las reuniones que se mantienen con el cliente durante el desarrollo "
    "de la consultoría/implantación. Su objetivo es que, por cada reunión, puedas:")
add_bullet("subir el acta o resumen de la reunión (.docx, .pdf o .txt);", "Subir el acta: ")
add_bullet("obtener automáticamente una síntesis de ~10 líneas mediante Inteligencia Artificial;", "Sintetizar: ")
add_bullet("extraer las tareas derivadas de cada reunión, con responsable y fecha límite;", "Extraer tareas: ")
add_bullet("hacer seguimiento y monitorización: completar, desestimar o reabrir tareas, conservando el histórico.", "Monitorizar: ")

# ============================================================
# 2. CONCEPTO CLAVE: CÓMO TRABAJA LA IA
# ============================================================
add_heading("2. Concepto clave: cómo trabaja la IA aquí", 1)
add_body(
    "Para que la IA sintetice las actas y extraiga tareas, la aplicación necesita conectarse "
    "a un servicio de Inteligencia Artificial. Esa conexión se realiza con una «API Key» "
    "(una credencial, similar a una contraseña, que da acceso al servicio). Existen DOS formas "
    "de trabajar:")

add_heading("2.1. Modo automático (con API Key)", 2)
add_body(
    "Configuras una API Key una sola vez. A partir de entonces, al pulsar «Sintetizar con IA», "
    "la aplicación llama sola al servicio y rellena el resumen y las tareas. Es el modo recomendado.")

add_heading("2.2. Modo manual (sin API Key — flujo de respaldo)", 2)
add_body(
    "Si no tienes API Key, la aplicación te muestra el «prompt» (el texto de instrucciones) para "
    "que lo copies, lo pegues en ChatGPT o Claude por tu cuenta, y después pegues de vuelta el "
    "resultado en la aplicación con el botón «Pegar resultado IA». Funciona igual; solo que el "
    "viaje a la IA lo haces tú manualmente.")

add_callout(
    "¿Por qué antes «no pasaba nada»?",
    "Si una reunión no funcionaba, no era porque los datos estuvieran dañados: simplemente no "
    "había una API Key configurada y el flujo manual no estaba visible. La versión actual corrige "
    "esto: la Configuración de la IA aparece ahora en la propia pestaña «Seguimiento», y los "
    "mensajes de error indican con claridad qué hacer.",
    title_color=ACCENT)

# ============================================================
# 3. NOVEDADES (TRAS LA ACTUALIZACIÓN)
# ============================================================
add_heading("3. Novedades (tras la actualización)", 1)
add_bullet("Flujo unificado por reunión: subir acta → resumen de ~10 líneas → tareas, en un solo paso.", "")
add_bullet("Configuración de IA accesible desde la propia pestaña «Seguimiento».", "")
add_bullet("Sistema de proveedores ABIERTO: puedes usar cualquier servicio de IA (incluidos gratuitos) sin tocar el código.", "")
add_bullet("Tareas con responsable, fecha límite y aviso de vencimiento.", "")
add_bullet("Gestión con histórico: Completar, Desestimar (con motivo) y Reabrir — sin perder el registro.", "")
add_bullet("Mensajes de error claros (clave vacía, modelo inexistente, cuota agotada, bloqueo del navegador).", "")

# ============================================================
# 4. CONFIGURAR LA IA (MODO AUTOMÁTICO)
# ============================================================
add_heading("4. Configurar la IA (modo automático)", 1)
add_body(
    "En la pestaña «Seguimiento», arriba, despliega el bloque «🔑 Configuración API (LLM)». "
    "Si no hay clave configurada, aparece ya abierto y resaltado.")

add_heading("4.1. Proveedores y modelo gratuito recomendado", 2)
add_body(
    "El selector «Proveedor» incluye opciones gratuitas (OpenRouter, Groq, Google Gemini) y de "
    "pago (OpenAI, Anthropic), además de «Personalizado» (cualquier servicio compatible). "
    "Para un uso gratuito y estable se recomienda:")
add_bullet("OpenRouter, con el modelo «openrouter/free» (alias estable que siempre apunta a un modelo gratuito disponible).", "OpenRouter: ")
add_bullet("Groq, con «llama-3.3-70b-versatile» (más rápido y con cuota amplia), si necesitas procesar varias actas seguidas.", "Groq: ")

add_callout(
    "Importante: los modelos gratuitos cambian a menudo",
    "El campo «Modelo» es editable. Si un modelo deja de existir (error 404) o se satura "
    "(error 429), simplemente escribe otro modelo gratuito en ese campo. Por eso «openrouter/free» "
    "es la opción más cómoda: no caduca.")

add_heading("4.2. Pasos para configurar", 2)
add_steps_table([
    ("1", "Pestaña «Seguimiento» → desplegar «🔑 Configuración API (LLM)»."),
    ("2", "Proveedor: elegir «OpenRouter (Gratis)» (o Groq)."),
    ("3", "Pulsar «Consigue tu API Key gratis» y crear la clave en la web del proveedor (gratis, normalmente sin tarjeta)."),
    ("4", "Pegar la clave en el campo «API Key»."),
    ("5", "Modelo: dejar «openrouter/free» (o elegir otro de la lista)."),
    ("6", "Pulsar «Probar conexión». Debe mostrar «✔ OK»."),
])

# ============================================================
# 5. SINTETIZAR UN ACTA
# ============================================================
add_heading("5. Sintetizar un acta de reunión", 1)
add_steps_table([
    ("1", "En la fila de la reunión, abrir «Acta / Notas»."),
    ("2", "En el bloque «Acta de la reunión», pulsar «Subir acta (.docx/.pdf/.txt)» y seleccionar el fichero."),
    ("3", "Pulsar «✨ Sintetizar con IA»."),
    ("4", "Esperar unos segundos: se rellenan el «Resumen IA (~10 líneas)» y las «Próximas tareas»."),
])
add_body(
    "Nota: los modelos gratuitos pueden tardar algo (es normal). Si una llamada falla por lentitud "
    "o saturación, la aplicación reintenta automáticamente una vez.", italic=True, color=GREY, size=10)

add_heading("5.1. Modo manual (sin API Key)", 2)
add_steps_table([
    ("1", "Pulsar «✨ Sintetizar con IA»: se copia el prompt al portapapeles y se abre una ventana con el texto."),
    ("2", "Pegar ese texto en ChatGPT o Claude (web) y enviar."),
    ("3", "Copiar la respuesta (en formato JSON) que devuelve la IA."),
    ("4", "En la aplicación, pulsar «Pegar resultado IA», pegar el JSON y aplicar."),
])

# ============================================================
# 6. GESTIÓN Y SEGUIMIENTO DE TAREAS
# ============================================================
add_heading("6. Gestión y seguimiento de tareas", 1)
add_body("Cada tarea (generada por IA o creada a mano) incluye:")
add_bullet("Título y descripción.", "")
add_bullet("Responsable (persona o rol).", "")
add_bullet("Fecha límite (con aviso de «Vencida» si pasa la fecha y no está cerrada).", "")
add_bullet("Estado: En curso, A revisar, Completada o Desestimada.", "")
add_bullet("Indicador «IA» si la propuso la Inteligencia Artificial.", "")

add_heading("6.1. Cerrar tareas conservando el histórico", 2)
add_bullet("se marca como «Completada» y se registra la fecha de cierre. Se puede añadir una nota (opcional).",
           "Completar: ")
add_bullet("para tareas que la IA propuso pero no son correctas. Exige un MOTIVO obligatorio, que queda en el histórico.",
           "Desestimar: ")
add_bullet("una tarea cerrada o desestimada vuelve a «En curso». El cambio también queda registrado.",
           "Reabrir: ")
add_body(
    "Ninguna acción borra la tarea: todas las tareas (completadas y desestimadas incluidas) se "
    "conservan y pueden consultarse filtrando por estado. Cada tarea guarda un «Historial» con "
    "todos los cambios (fecha, acción y nota).")

add_callout(
    "Panel de seguimiento de tareas",
    "Al final de la pestaña «Seguimiento» hay un panel que agrupa TODAS las tareas de todas las "
    "reuniones, con filtros por estado (Todas, En curso, A revisar, Completada, Desestimada, "
    "Vencida) y contadores. Desde ahí puedes Completar, Desestimar o Reabrir, y ver responsable y "
    "fecha límite de cada tarea.",
    title_color=ACCENT)

# ============================================================
# 7. PROBAR SOBRE UN PROYECTO EXISTENTE
# ============================================================
add_heading("7. Probar sobre un proyecto existente (p. ej. PRJ-006-REDMAN)", 1)
add_body(
    "Los datos de las reuniones se guardan localmente en el navegador. Si un proyecto ya tiene "
    "reuniones creadas, NO se pierden: la actualización añade los campos nuevos de forma no "
    "destructiva.")
add_steps_table([
    ("1", "Abrir el proyecto → pestaña «Seguimiento». Las reuniones existentes siguen ahí."),
    ("2", "Configurar la API Key (sección 4) si aún no está."),
    ("3", "En una reunión, subir el acta y pulsar «✨ Sintetizar con IA»."),
    ("4", "Revisar el resumen y las tareas; asignar responsables y fechas."),
    ("5", "Probar Completar / Desestimar / Reabrir y revisar el histórico."),
    ("6", "Recargar la página y comprobar que todo se ha guardado."),
])
add_body(
    "Nota: si esas reuniones «no funcionaban» antes, era por falta de API Key (no por datos "
    "corruptos). Una vez configurada, funcionan con normalidad.", italic=True, color=GREY, size=10)

# ============================================================
# 8. PREGUNTAS FRECUENTES Y SOLUCIÓN DE PROBLEMAS
# ============================================================
add_heading("8. Preguntas frecuentes y solución de problemas", 1)

faqs = [
    ("Al pulsar «Sintetizar con IA» no pasa nada.",
     "Comprueba que hay una API Key configurada y que «Probar conexión» da OK. Sin clave, se "
     "abrirá la ventana del modo manual para copiar el prompt."),
    ("«Probar conexión» da error de conexión / CORS.",
     "Algunos navegadores (Safari) bloquean las llamadas a la IA al abrir el archivo con file://. "
     "Usa un proveedor gratuito (OpenRouter/Groq/Gemini), o sirve la aplicación por http://localhost, "
     "o usa el flujo manual."),
    ("Error 404: «No endpoints found for …».",
     "El nombre del modelo ya no existe. Escribe otro modelo en el campo «Modelo» "
     "(por ejemplo «openrouter/free»)."),
    ("Error 429: cuota agotada.",
     "El modelo gratuito está saturado. Cambia a otro modelo :free, o usa el proveedor Groq "
     "(cuota más amplia)."),
    ("La IA devuelve textos de ejemplo en vez de contenido real.",
     "Suele ocurrir con modelos poco potentes o actas con poco texto. La aplicación filtra y limpia "
     "esos textos; si persiste, usa un modelo mejor o revisa que el acta tenga contenido."),
    ("El resumen tarda mucho.",
     "Es normal con modelos gratuitos. La aplicación espera hasta 90 segundos y reintenta una vez. "
     "Para mayor velocidad, usa Groq."),
]
for q, a in faqs:
    p = doc.add_paragraph()
    rq = p.add_run("P: " + q)
    rq.bold = True
    rq.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(1)
    pa = doc.add_paragraph()
    ra = pa.add_run("R: " + a)
    ra.font.size = Pt(10.5)
    pa.paragraph_format.space_after = Pt(6)

# ============================================================
# 9. GLOSARIO
# ============================================================
add_heading("9. Glosario", 1)
glos = [
    ("API Key", "Credencial que autoriza a la aplicación a usar un servicio de IA."),
    ("Prompt", "Texto de instrucciones que se envía a la IA para que genere el resultado."),
    ("Modelo (LLM)", "El motor de IA concreto que procesa el texto (p. ej. Llama, Gemini, GPT)."),
    ("Síntesis", "Resumen breve (~10 líneas) generado por la IA a partir del acta."),
    ("CORS", "Mecanismo de seguridad del navegador que puede bloquear llamadas externas con file://."),
    ("Fallback (manual)", "Flujo de respaldo: copiar el prompt, usar ChatGPT/Claude y pegar el resultado."),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light List Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Término"
hdr[1].text = "Significado"
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True
for term, mean in glos:
    cells = tbl.add_row().cells
    cells[0].text = term
    cells[1].text = mean

# Pie de página
set_footer()

doc.save(OUT)
print("Manual generado en:", OUT)
